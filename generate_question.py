import re
import sys
from docx import Document

def classify(text):
    """Assign subject, unit, topic based on keywords in the question."""
    t = text.lower()
    if "percent" in t or "%" in t:
        return ("Quantitative Math", "Numbers and Operations", "Fractions, Decimals, & Percents")
    if "area" in t or "circle" in t or "square" in t:
        return ("Quantitative Math", "Geometry and Measurement", "Circles (Area, circumference)")
    if "sequence" in t or "repeat" in t:
        return ("Quantitative Math", "Data Analysis & Probability", "Counting & Arrangement Problems")
    if "$x$" in t or "$n$" in t or "equation" in t:
        return ("Quantitative Math", "Algebra", "Interpreting Variables")
    return ("Quantitative Math", "Problem Solving", "Problem Solving")

def extract_question(block):
    """Extract stem, options, and image URLs from a question block."""
    lines = block.splitlines()
    if re.match(r'^\d+\.\s', lines[0]):
        lines[0] = re.sub(r'^\d+\.\s*', '', lines[0])
    text = "\n".join(lines)

    img_urls = re.findall(r'!\[.*?\]\((.*?)\)', text)
    parts = re.split(r'\n(?=\([A-E]\))', text)
    stem = parts[0].strip()
    options = [p.strip() for p in parts[1:]]
    norm_opts = []
    for opt in options:
        m = re.match(r'^\(([A-E])\)\s*(.*)', opt, re.S)
        if m:
            norm_opts.append((m.group(1), m.group(2).strip()))
    return stem, norm_opts, img_urls

def main(input_md, output_docx):
    # Read markdown file
    with open(input_md, 'r', encoding='utf-8') as f:
        md = f.read()

    # Split into question blocks
    blocks = re.split(r'(?m)^(?=\d+\.\s)', md.strip())
    blocks = [b.strip() for b in blocks if b.strip()]

    # Create Word document
    doc = Document()
    doc.add_heading("ML Official T1 – 25 Questions (Formatted Output)", 0)
    doc.add_paragraph("Reformatted into @tag schema from provided dataset. Images are preserved via URL references.")

    for idx, block in enumerate(blocks, start=1):
        stem, options, img_urls = extract_question(block)
        subject, unit, topic = classify(stem)
        difficulty = "easy" if len(stem) < 120 else "moderate"

        doc.add_paragraph(f"@title Q{idx}: Auto Reformatted")
        doc.add_paragraph("@description Reformatted from provided dataset; images preserved via URL where applicable.")
        
        qtext = stem
        if img_urls:
            qtext += "\n\n" + "\n".join(f"[Image: {u}]" for u in img_urls)
        doc.add_paragraph("@question " + qtext)

        doc.add_paragraph("@instruction Choose the correct option.")
        doc.add_paragraph(f"@difficulty {difficulty}")
        doc.add_paragraph(f"@Order {idx}")

        for letter, content in options:
            doc.add_paragraph(f"@option {content}")

        doc.add_paragraph("@explanation ")
        doc.add_paragraph("Explanation not provided in the source.")

        doc.add_paragraph(f"@subject {subject}")
        doc.add_paragraph(f"@unit {unit}")
        doc.add_paragraph(f"@topic {topic}")
        doc.add_paragraph("@plusmarks 1")
        doc.add_paragraph("")

    # Save output
    doc.save(output_docx)
    print(f"Saved formatted questions to {output_docx}")

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python generate_questions.py input.md output.docx")
    else:
        main(sys.argv[1], sys.argv[2])
